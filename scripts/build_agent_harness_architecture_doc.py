"""Build the FYP agent-harness architecture explanation DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "FYP_Agent_Harness_Architecture_Guide_2026-06-09.docx"
FIGURE = ROOT / "docs" / "architecture" / "fyp_quant_integrated_agentic_stack.png"

NAVY = "0F172A"
BLUE = "2F80C5"
TEAL = "2494A3"
GREEN = "2C9A67"
PURPLE = "8A63C8"
ORANGE = "C6812D"
RED = "C85C74"
SLATE = "334155"
LIGHT_BLUE = "EEF6FF"
LIGHT_TEAL = "EFFAFB"
LIGHT_PURPLE = "F7F0FF"
LIGHT_ORANGE = "FFF5E8"
LIGHT_GREEN = "F2FAF5"
LIGHT_RED = "FFF0F3"
LIGHT_SLATE = "F8FAFC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "CBD5E1", size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, start: int = 120, bottom: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_run(run, size: int = 10, bold: bool = False, color: str = SLATE) -> None:
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", size: int = 10, color: str = SLATE, bold: bool = False, space_after: int = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    style_run(run, size=size, color=color, bold=bold)
    return p


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=18 if level == 1 else 13, bold=True, color=NAVY if level == 1 else BLUE)


def add_bullets(doc: Document, items: list[str], size: int = 9) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run(item)
        style_run(run, size=size, color=SLATE)


def add_numbered(doc: Document, items: list[str], size: int = 9) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        run = p.add_run(item)
        style_run(run, size=size, color=SLATE)


def add_step_blocks(doc: Document, items: list[str]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [0.65, 8.35])
    for idx, item in enumerate(items, start=1):
        row = table.add_row()
        number_cell, text_cell = row.cells
        for cell in row.cells:
            set_cell_border(cell, "CBD5E1", "8")
            set_cell_margins(cell, 100, 130, 100, 130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(number_cell, NAVY)
        set_cell_shading(text_cell, "FFFFFF")
        p = number_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(idx))
        style_run(run, size=12, bold=True, color="FFFFFF")
        run2 = text_cell.paragraphs[0].add_run(item)
        style_run(run2, size=9, color=SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE, border: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [9.0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, border, "10")
    set_cell_margins(cell, 130, 170, 130, 170)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    style_run(run, size=10, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run2 = p2.add_run(body)
    style_run(run2, size=9, color=SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_two_column_cards(doc: Document, cards: list[tuple[str, str, list[str], str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [4.45, 4.45])
    for idx in range(0, len(cards), 2):
        row = table.add_row()
        for col in range(2):
            cell = row.cells[col]
            set_cell_margins(cell, 125, 150, 125, 150)
            set_cell_border(cell, "CBD5E1", "8")
            if idx + col >= len(cards):
                set_cell_shading(cell, "FFFFFF")
                continue
            title, subtitle, bullets, fill, border = cards[idx + col]
            set_cell_shading(cell, fill)
            set_cell_border(cell, border, "10")
            p = cell.paragraphs[0]
            run = p.add_run(title)
            style_run(run, size=10, bold=True, color=NAVY)
            p2 = cell.add_paragraph()
            run2 = p2.add_run(subtitle)
            style_run(run2, size=8, bold=True, color=border)
            for bullet in bullets:
                bp = cell.add_paragraph()
                bp.paragraph_format.left_indent = Inches(0.15)
                bp.paragraph_format.first_line_indent = Inches(-0.1)
                bp.paragraph_format.space_after = Pt(1)
                br = bp.add_run(f"- {bullet}")
                style_run(br, size=8, color=SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_scenario_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [2.2, 6.8])
    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("Stage", "What happens in practice")):
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY, "8")
        set_cell_margins(cell, 120, 130, 120, 130)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        style_run(run, size=9, bold=True, color="FFFFFF")
    for stage, detail in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_border(cell)
            set_cell_margins(cell, 110, 130, 110, 130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_shading(cells[0], LIGHT_SLATE)
        set_cell_shading(cells[1], "FFFFFF")
        run = cells[0].paragraphs[0].add_run(stage)
        style_run(run, size=8, bold=True, color=NAVY)
        run2 = cells[1].paragraphs[0].add_run(detail)
        style_run(run2, size=8, color=SLATE)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(SLATE)
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(9)


def build_doc() -> None:
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("fyp_quant Repo-Native Agent Harness")
    style_run(run, size=25, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run("Architecture explanation, presentation script, and practical workflow guide")
    style_run(run, size=11, bold=True, color=BLUE)

    add_callout(
        doc,
        "One-sentence explanation",
        "I use Codex as a repo-native research operating system: the repo gives the agent durable memory, scoped tasks, context-isolated reviewers, executable checks, protected evidence, and recoverable handoffs.",
        fill=LIGHT_TEAL,
        border=TEAL,
    )

    add_two_column_cards(
        doc,
        [
            (
                "For presentation",
                "What to say on stage",
                ["Use the diagram as the main visual anchor.", "Explain the harness as a research operating system."],
                LIGHT_BLUE,
                BLUE,
            ),
            (
                "For your own understanding",
                "How to reason about the repo",
                ["Separate repo memory, task routing, subagents, checks, and evidence.", "Use the scenarios as your mental model."],
                LIGHT_GREEN,
                GREEN,
            ),
            (
                "For workflow demos",
                "What to actually show",
                ["Run agent-status, agent-start, then make agent-check.", "Show how handoff/dashboard files recover the next session."],
                LIGHT_PURPLE,
                PURPLE,
            ),
            (
                "For research integrity",
                "Why this is more than productivity",
                ["Raw TC1 outputs are protected.", "Scoring corrections become derived sidecars with a logged audit trail."],
                LIGHT_ORANGE,
                ORANGE,
            ),
        ],
    )

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(FIGURE), width=Inches(9.2))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run("Figure 1. Integrated agentic stack: folder hierarchy, harness lifecycle, subagents, checks, benchmark pipeline, and evidence contract.")
    style_run(run, size=8, color=SLATE)

    doc.add_page_break()

    add_heading(doc, "How to read the architecture", 1)
    add_two_column_cards(
        doc,
        [
            (
                "1. Folder hierarchy",
                "Where the repo stores responsibility",
                [
                    "Agent OS files store instructions, state, task packets, skills, subagents, hooks, and CI.",
                    "Research core files run benchmarks, model loading, analysis, judging, and SLURM helpers.",
                    "Evidence files store raw outputs, derived sidecars, analysis, report, and visuals.",
                ],
                LIGHT_BLUE,
                BLUE,
            ),
            (
                "2. Agent harness control plane",
                "How a fresh session becomes safe and useful",
                [
                    "The agent starts from AGENTS.md and docs/PROJECT_LOG.md, not from chat memory.",
                    "agent-start loads a small task packet and recommends a focused subagent.",
                    "make agent-check is the finish gate before handoff or commit.",
                ],
                LIGHT_TEAL,
                TEAL,
            ),
            (
                "3. Multi-agent layer",
                "How context isolation works",
                [
                    "The main Codex session owns edits, final judgement, and project-log updates.",
                    "Subagents audit one slice: report, artifacts, TC1 ops, judge validation, or meetup story.",
                    "Skills are loaded only when relevant, reducing context pressure.",
                ],
                LIGHT_PURPLE,
                PURPLE,
            ),
            (
                "4. Evidence contract",
                "Why the system protects research integrity",
                [
                    "raw.jsonl and summary.json are treated as immutable TC1-original evidence.",
                    "Scoring corrections write sidecars such as scores.v2.* and scores.judge.*.",
                    "Generated handoff/dashboard artifacts help the next agent recover without exposing raw HarmBench text.",
                ],
                LIGHT_ORANGE,
                ORANGE,
            ),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "What the harness actually does", 1)
    add_callout(
        doc,
        "The mental model",
        "The harness converts fragile prompt instructions into repo-native mechanisms: source-of-truth documents, task packets, machine-readable artifact policy, custom subagents, hooks, generated handoffs, and executable checks.",
        fill=LIGHT_GREEN,
        border=GREEN,
    )
    add_two_column_cards(
        doc,
        [
            (
                "Orientation",
                "What every agent reads first",
                ["AGENTS.md / CLAUDE.md: operating law.", "docs/PROJECT_LOG.md: current truth.", "docs/HANDOFF.md: short bridge, never source of truth."],
                LIGHT_SLATE,
                BLUE,
            ),
            (
                "Task routing",
                "How the agent avoids loading everything",
                ["python fyp_cli.py agent-status", "python fyp_cli.py agent-start --task T21 --agent fyp-report-auditor", "docs/agent_tasks/ keeps the task bounded."],
                LIGHT_SLATE,
                GREEN,
            ),
            (
                "Context isolation",
                "How review stays focused",
                ["fyp-report-auditor for report/log drift.", "fyp-artifact-guardian for raw evidence and redaction.", "fyp-tc1-ops, fyp-judge-reviewer, fyp-meetup-story for specialist checks."],
                LIGHT_SLATE,
                PURPLE,
            ),
            (
                "Verification",
                "How the repo says done",
                ["make agent-check checks docs sync, project-log discipline, immutable artifacts, stale text, redaction, report freshness, whitespace, and tests.", "make harness-eval proves the harness catches bad states."],
                LIGHT_SLATE,
                TEAL,
            ),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "Workflow examples", 1)
    add_para(
        doc,
        "These scenarios are the practical story you can present: a human gives intent, the repo narrows context, Codex does scoped work, specialist agents audit risky areas, and the finish gate turns the work into evidence.",
        size=9,
    )

    add_scenario_table(
        doc,
        "Scenario 1: Starting a new Codex session for FYP report work",
        [
            ("Human intent", "You say: 'Strengthen Chapter 6/7 results discussion and keep the judge-vs-v2 disagreement central.'"),
            ("Agent entry", "Codex reads AGENTS.md and docs/PROJECT_LOG.md, then runs python fyp_cli.py agent-status to understand live repo state."),
            ("Task packet", "Codex runs python fyp_cli.py agent-start --task T21 --agent fyp-report-auditor. This loads a bounded packet instead of the entire repo story."),
            ("Specialist audit", "The fyp-report-auditor reviews report source, PROJECT_LOG, judge_agreement outputs, and stale current-facing claims. It returns findings, not random edits."),
            ("Implementation", "The main Codex session edits scripts/build_fyp_report_v3.js if report-worthy, regenerates the report with make report, and updates docs/PROJECT_LOG.md."),
            ("Finish gate", "Codex runs make agent-check. If it passes, generated handoff/dashboard files are refreshed so another session can continue."),
        ],
    )

    doc.add_page_break()

    add_scenario_table(
        doc,
        "Scenario 2: Protecting evidence after a scoring correction",
        [
            ("Problem", "The scoring layer is wrong, but the model outputs are still valid. This was the key lesson from the refusal parser and HarmBench judge work."),
            ("Harness rule", "The raw TC1 outputs remain immutable: results/*/*/raw.jsonl and summary.json are not overwritten."),
            ("Safe path", "New scoring writes sidecars: scores.v2.*, summary.v2.*, scores.judge.*, and summary.judge.*."),
            ("Subagent use", "The fyp-artifact-guardian checks immutability and redaction. The fyp-judge-reviewer checks judge sidecars and agreement analysis."),
            ("Human explanation", "In the presentation: 'I did not ask the model to remember what happened. I made the repo remember the evidence trail.'"),
            ("Verification", "make agent-check confirms the immutable manifest, redaction scan, stale-text scan, report freshness, and tests."),
        ],
    )

    doc.add_page_break()

    add_scenario_table(
        doc,
        "Scenario 3: Running or reviewing TC1 cluster work",
        [
            ("Human boundary", "You handle identity, VPN, MFA, and any private authentication. The agent should not pretend that these can be automated away."),
            ("Agent role", "Codex prepares and reviews sbatch scripts, TC1 runbook steps, offline-mode variables, and safe submission commands."),
            ("TC1 policy", "The fyp-tc1-ops subagent checks that real compute uses sbatch, not head-node Python or interactive srun."),
            ("Artifacts", "Jobs write results into the expected results/ structure. Logs and summaries can be inspected later without exposing raw HarmBench content."),
            ("Recovery", "docs/TC1_AGENT_CHECKLIST.md and docs/HANDOFF.md let a future agent continue from the last verified state."),
        ],
    )

    doc.add_page_break()

    add_scenario_table(
        doc,
        "Scenario 4: Preparing the Codex meetup explanation",
        [
            ("Goal", "You want to explain not just the FYP result, but how you used Codex as a research partner."),
            ("Task packet", "Use python fyp_cli.py agent-start --task codex-meetup-prep --agent fyp-meetup-story."),
            ("Story spine", "The fyp-meetup-story agent turns repo mechanisms into a human explanation: repo memory, task packets, subagents, hooks, checks, immutable evidence, and handoff recovery."),
            ("USP", "The strongest line: 'Memory is not magic; make the repo remember.'"),
            ("Demo", "Show Figure 1, then run agent-status and agent-start. Finish by showing make agent-check as the evidence-producing gate."),
        ],
    )

    doc.add_page_break()

    add_heading(doc, "Presentation script", 1)
    add_step_blocks(
        doc,
        [
            "Start with the pain: long research projects outlive a single chat window, so chat memory alone is not reliable enough.",
            "Show the folder hierarchy: the repo contains both the benchmark framework and an agent operating layer.",
            "Explain the control plane: AGENTS.md gives law, PROJECT_LOG gives truth, agent-start gives task scope, subagents isolate audits, and make agent-check proves the state.",
            "Show the evidence contract: raw model outputs are immutable; scoring corrections become sidecars; report text is regenerated from source.",
            "Give the scoring-correction example: the model outputs were still valid, but the scoring layer was wrong, so the harness preserved raw evidence while deriving new scores.",
            "Close with the USP: I use Codex less like autocomplete and more like a repo-native research operating system.",
        ],
    )

    add_heading(doc, "What to say when someone asks 'why not just prompt better?'", 1)
    add_callout(
        doc,
        "Answer",
        "Prompting tells the agent what to do once. A harness makes the repo keep enforcing the important rules: where truth lives, what must not be changed, which context to load, which checks prove completion, and how the next session resumes.",
        fill=LIGHT_RED,
        border=RED,
    )

    doc.add_page_break()

    add_heading(doc, "Practical commands to remember", 1)
    add_bullets(
        doc,
        [
            "Fresh session: python fyp_cli.py agent-status",
            "Task-specific startup: python fyp_cli.py agent-start --task T21 --agent fyp-report-auditor",
            "Regenerate visual pack: make architecture-diagram",
            "Refresh handoff/dashboard: make agent-handoff && make agent-dashboard && make agent-tc1-checklist",
            "Final finish gate: make agent-check",
        ],
        size=9,
    )

    add_heading(doc, "The one-minute story", 1)
    add_para(
        doc,
        "In my FYP, Codex is not just writing code. I turned the repo into an agent harness. The repo tells Codex what the current truth is, gives it bounded task packets, lets it call specialist subagents for audits, protects raw experiment evidence, and forces a finish gate before I trust the result. This matters because my project had a real scoring correction: the model outputs stayed valid, but the scoring layer changed. Because the harness preserved raw outputs and wrote sidecars, the research story stayed auditable instead of becoming a chat-memory mess.",
        size=10,
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("fyp_quant agent harness architecture guide - generated from repo state on 2026-06-09")
    style_run(run, size=8, color="64748B")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build_doc()
